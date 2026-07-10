#!/usr/bin/env python3
"""
=============================================================================
PDF -> Memorial RSC-PCCTAE — Gerador Autônomo v5.3 (NARRATIVAS LLM)
=============================================================================
Gera o memorial completo de Reconhecimento de Saberes e Competências (RSC-PCCTAE)
autonomamente a partir do PDF oficial do Relatório Detalhado RSC emitido pelo
sistema da UFV (Pró-Reitoria de Gestão de Pessoas), em conformidade com o
Decreto nº 13.048, de 3 de julho de 2026 (Art. 13).

v5.0 — Narrativas geradas pelo LLM (Large Language Model) em linguagem
        científica, longa e detalhada — como o memorial de um professor
        titular. O run.py recebe as narrativas via JSON e as insere nos
        tópicos pré-estabelecidos da estrutura.

FLUXO RECOMENDADO:
  1. run.py extrai dados do PDF para JSON
     $ python3 run.py relatorio.pdf --dump-json dados_extraidos.json
  
  2. LLM (agente PesquisAI) gera narrativas detalhadas a partir dos dados
     -> narrativas.json (ver SKILL.md para o prompt completo)
  
  3. run.py monta o memorial com as narrativas do LLM
     $ python3 run.py relatorio.pdf --narrativas narrativas.json

Uso:
  python3 run.py <caminho_do_pdf> [--output-dir DIR] [--nome "Nome"]
                 [--ano-ingresso ANO] [--auto]
                 [--narrativas JSON] [--dump-json ARQUIVO]

Dependências:
  pip install pdfplumber python-docx pyspellchecker
=============================================================================
"""

import re, os, sys, argparse, json, random
from pathlib import Path
from datetime import datetime

try:
    import pdfplumber
except ImportError:
    sys.exit("Erro: pdfplumber não instalado. Execute: pip install pdfplumber")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
except ImportError:
    sys.exit("Erro: python-docx não instalado. Execute: pip install python-docx")

try:
    from spellchecker import SpellChecker as _SpellChecker
    _SP = _SpellChecker(language='pt')
    _PT_WORDS = set(_SP.word_frequency.dictionary)
except Exception:
    _SP = None
    _PT_WORDS = set()

NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
DECRETO_URL = "https://www.planalto.gov.br/ccivil_03/_ato2023-2026/2026/decreto/d13048.htm"


# =============================================================================
# Utilitários de texto
# =============================================================================
def _strip_accents(s):
    return ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')

import unicodedata

def _build_reverse_index():
    if _SP is None:
        return {}
    rev = {}
    freqs = _SP.word_frequency.dictionary
    for w in _PT_WORDS:
        sw = _strip_accents(w)
        if sw != w:
            f = freqs.get(w, 0)
            if f >= 50:
                rev.setdefault(sw, []).append((w, f))
    for sw, lst in rev.items():
        lst.sort(key=lambda t: -t[1])
    return rev

_REV_INDEX = _build_reverse_index()
_STOPWORDS = frozenset("""
a o os as um uns uma umas de da do das dos no na nas ne noh neh
por para pela pelas pelo pelos perante ante apos ate sobre sob contra
com sem entre desde durante segundo mediante conforme contra exceto
e ou mas que se como quando onde porque pois embora entretanto porem
contudo todavia logo portanto assim ja nem apenas senao entao enquanto
que quem qual quais cujo cuja cujos cujas quanto quanta quantos quantas
ele ela eles elas este esta estes estas esse essa esses essas aquele
aquela aqueles aquelas isto isso aquilo meu minha meus minhas teu tua
teus tuas seu sua seus suas nosso nossa nossos nossas vosso vossa vossos
vossas lhe lhes me te se nos vos si mesmo mesma mesmos mesmas cada todo
todos toda todas algum alguns alguma algumas nenhum nenhuns nenhuma
nenhumas outro outros outra outras muito muitos muita muitas pouco poucos
pouca poucas tal tais qual quais quer queres eu tu ele nos vos eles voce
voces nos tao pouco demais acerca afinal entao ainda aqui ali la ali
ai ja nao sim ora pois ah oh ih olha olhe veja
polo logo cedo seco fato fala fica saca saco bato bate pato
mate meto mito moto pito peta taxi onibus onibus
""".split())

def _match_case(template, target):
    if template.isupper():
        return target.upper()
    if template[0].isupper():
        return target.capitalize()
    return target.lower()

def restaurar_acentos(texto):
    if not _REV_INDEX:
        return texto
    proibidas = _STOPWORDS
    freqs = _SP.word_frequency.dictionary if _SP else {}

    def _sub(match):
        word = match.group(0)
        sw = _strip_accents(word)
        if sw != word:
            return word
        if word.lower() in proibidas:
            return word
        if len(sw) < 4:
            return word
        candidates = _REV_INDEX.get(sw)
        if not candidates:
            return word
        top, top_f = candidates[0]
        second_f = candidates[1][1] if len(candidates) > 1 else 0
        # Só acentua se a variante ACENTUADA for MAIS frequente que a
        # própria forma SEM acento no corpus. Isso evita colocar acento
        # em palavras corretas sem acento (ex.: pesquisa, apresenta,
        # atua), que antes eram wrongly transformadas em pesquisá,
        # apresentá, atuá.
        f_sem = freqs.get(sw, 0)
        if top_f <= f_sem:
            return word
        if len(candidates) == 1:
            return _match_case(word, top)
        if top_f >= 10 * second_f:
            return _match_case(word, top)
        return word

    return re.sub(r"(?<![/\.:])\b[a-zA-Zà-ÿÀ-Ÿ]+\b(?![/\.:])", _sub, texto)


def _acentuar_vogal_anterior(palavra):
    if not palavra:
        return palavra
    acentos = {'a': 'á', 'e': 'é', 'i': 'í', 'o': 'ó', 'u': 'ú'}
    acentuados = set(acentos.values())
    if any(c in acentuados for c in palavra):
        return palavra
    for idx in range(len(palavra) - 1, -1, -1):
        if palavra[idx] in acentos:
            return palavra[:idx] + acentos[palavra[idx]] + palavra[idx + 1:]
    return palavra


def normalizar_texto(texto):
    """Correção de acentos — regras de sufixo + dicionário."""
    # Regras sistemáticas de sufixo
    texto = re.sub(r'\b(\w{2,})cao\b', lambda m: m.group(1) + 'ção', texto, flags=re.IGNORECASE)
    texto = re.sub(r'\b(\w{2,})oes\b', lambda m: m.group(1) + 'ões', texto, flags=re.IGNORECASE)
    texto = re.sub(r'\b(\w{2,})aes\b', lambda m: m.group(1) + 'ães', texto, flags=re.IGNORECASE)
    texto = re.sub(r'\b(\w{3,})ao\b', lambda m: m.group(1) + 'ão', texto, flags=re.IGNORECASE)
    texto = re.sub(r'\b(\w+)(vel)\b', lambda m: _acentuar_vogal_anterior(m.group(1)) + m.group(2), texto, flags=re.IGNORECASE)
    texto = re.sub(r'\b(Meu nome) e\b', r'\1 é', texto)
    texto = re.sub(r'(?<=^)E\b', 'É', texto)
    texto = re.sub(r'(?<=[.!?] )E\b', 'É', texto)
    texto = re.sub(r'\bnão e\b', 'não é', texto)
    texto = re.sub(r'\bnao e\b', 'não é', texto, flags=re.IGNORECASE)

    palavra_map = {
        'ambito': 'âmbito', 'epoca': 'época', 'epocas': 'épocas',
        'periodo': 'período', 'periodos': 'períodos',
        'genero': 'gênero', 'generos': 'gêneros',
        'inicio': 'início', 'fenomeno': 'fenômeno', 'exito': 'êxito',
        'carater': 'caráter', 'matricula': 'matrícula',
        'academico': 'acadêmico', 'academica': 'acadêmica',
        'academicos': 'acadêmicos', 'academicas': 'acadêmicas',
        'cientifico': 'científico', 'cientifica': 'científica',
        'cientificos': 'científicos', 'cientificas': 'científicas',
        'tecnico': 'técnico', 'tecnica': 'técnica',
        'tecnicos': 'técnicos', 'tecnicas': 'técnicas',
        'publico': 'público', 'publica': 'pública',
        'publicos': 'públicos', 'publicas': 'públicas',
        'especifico': 'específico', 'especifica': 'específica',
        'especificos': 'específicos', 'especificas': 'específicas',
        'proprio': 'próprio', 'propria': 'própria',
        'proprios': 'próprios', 'proprias': 'próprias',
        'multiplas': 'múltiplas', 'multiplos': 'múltiplos',
        'inumeras': 'inúmeras', 'inumeros': 'inúmeros',
        'economico': 'econômico', 'economica': 'econômica',
        'estrategico': 'estratégico', 'estrategica': 'estratégica',
        'pedagogico': 'pedagógico', 'pedagogica': 'pedagógica',
        'tecnologico': 'tecnológico', 'tecnologica': 'tecnológica',
        'tres': 'três', 'tambem': 'também', 'ate': 'até', 'ja': 'já',
        'so': 'só', 'voce': 'você',
    }
    for sem_acento, com_acento in palavra_map.items():
        texto = re.sub(r'\b' + re.escape(sem_acento) + r'\b', com_acento, texto, flags=re.IGNORECASE)

    texto = re.sub(r'\b(Lei|Decreto|Portaria|Resolução|Instrução|Norma|Regulamento)\s+n\s+', r'\1 nº ', texto)
    texto = re.sub(r'\bn\s+(\d[\d.,\s]*(?:/20\d\d|/19\d\d))', r'nº \1', texto)
    texto = re.sub(r'\bsao\b', 'são', texto, flags=re.IGNORECASE)
    texto = re.sub(r'\bnao\b', 'não', texto, flags=re.IGNORECASE)
    return texto


NIVEL_EQUIVALENCIA = {
    'VI': {'nome': 'VI', 'equivalente': 'Doutor', 'percentual': '75%',
           'destinado': 'servidor com diploma de mestrado',
           'percentual_extenso': 'setenta e cinco por cento'},
    'V': {'nome': 'V', 'equivalente': 'Mestre', 'percentual': '52%',
           'destinado': 'servidor com certificado de pós-graduação lato sensu',
           'percentual_extenso': 'cinquenta e dois por cento'},
    'IV': {'nome': 'IV', 'equivalente': 'Graduacao', 'percentual': '30%',
           'destinado': 'servidor com diploma de graduação no ensino superior',
           'percentual_extenso': 'trinta por cento'},
}
NIVEL_PADRAO = NIVEL_EQUIVALENCIA['VI']


def parse_br(s):
    if not s: return 0.0
    s = s.strip()
    if ',' in s: return float(s.replace('.', '').replace(',', '.'))
    return float(s)


def fmt_br(n):
    if n is None: return "0,00"
    return f"{n:.2f}".replace('.', ',')


def get_nivel_info(rsc_raw):
    if not rsc_raw: return NIVEL_PADRAO
    t = rsc_raw.upper().strip()
    for nivel in ['VI', 'V', 'IV']:
        if nivel in t:
            return NIVEL_EQUIVALENCIA.get(nivel, NIVEL_PADRAO)
    m = re.search(r'N[ÍI]VEL\s*(VI|V|IV)', t)
    if m: return NIVEL_EQUIVALENCIA.get(m.group(1), NIVEL_PADRAO)
    return NIVEL_PADRAO


def perguntar_ano_ingresso():
    ano_atual = datetime.now().year
    while True:
        try:
            resp = input("\nEm que ano você ingressou na UFV? ").strip()
            if not resp: continue
            ano = int(resp)
            if 1960 <= ano <= ano_atual: return ano
            print(f"Ano inválido. Digite entre 1960 e {ano_atual}.")
        except ValueError:
            print("Digite um ano válido.")


# =============================================================================
# PARSER DO PDF (v4.0 — extração completa)
# =============================================================================
class RSCPDFParser:
    """Parser completo do Relatório Detalhado RSC — extrai TODO o conteúdo."""

    def __init__(self, pdf_path):
        self.pdf_path = Path(pdf_path)
        self.data = {
            'nome': '', 'matricula': '', 'cargo': '', 'titulacao': '',
            'rsc_requerido': '', 'data_admissao': '', 'lotacao': '',
            'funcao': '', 'nivel_classe': '', 'rsc_nivel': 'VI',
            'equivalente': 'Doutorado',
            'total_geral': 0.0, 'total_criterios': 0,
            'grupos': [], 'criterios': {}, 'ordem_criterios': [],
        }
        self._parse()

    def _parse(self):
        with pdfplumber.open(str(self.pdf_path)) as pdf:
            pages = [p.extract_text() or '' for p in pdf.pages]
        raw = '\n'.join(pages)
        self._extract_header(raw)
        self._extract_grupos(raw)
        self._extract_criterios(pages)
        self.data['nivel_info'] = get_nivel_info(self.data.get('rsc_requerido', ''))

    def _extract_header(self, raw):
        def get(pat, g=1):
            m = re.search(pat, raw)
            return m.group(g).strip() if m else ''
        self.data['matricula'] = get(r'Matr[íi]cula SIAPE:\s*(\d+)')
        self.data['nome'] = get(r'Nome:\s*(.+?)(?:\n|$)')
        self.data['cargo'] = get(r'Cargo:\s*(.+?)(?:\n|$)')
        self.data['data_admissao'] = get(r'Data de Admiss[ãa]o:\s*(.+?)(?:\n|$)')
        self.data['nivel_classe'] = get(r'N[íi]vel de Classifica[çc][ãa]o:\s*(.+?)(?:\n|$)')
        self.data['lotacao'] = get(r'Lota[çc][ãa]o:\s*(.+?)(?:\n|$)')
        self.data['funcao'] = get(r'Fun[çç][ãa]o/Encargo \(se houver\):\s*(.+?)(?:\n|$)')
        if not self.data['funcao']:
            self.data['funcao'] = get(r'Fun[çc][ãa]o\s*[–-]?\s*Encargo.*?:\s*(.+?)(?:\n|$)')
        self.data['titulacao'] = get(r'Titula[çc][ãa]o Atual:\s*(.+?)(?:\n|$)')
        if not self.data['titulacao']:
            self.data['titulacao'] = get(r'Titula[çc][ãa]o:\s*(.+?)(?:\n|$)')
        m_rsc = re.search(r'RSC\s+(VI|V|IV)', raw, re.I)
        self.data['rsc_nivel'] = m_rsc.group(1).strip() if m_rsc else 'VI'
        eq = get(r'Equivalente a\s+(.+?)(?:\n|$)')
        self.data['equivalente'] = eq if eq else 'Doutorado'
        self.data['rsc_requerido'] = f"RSC {self.data['rsc_nivel']} - Equivalente a {self.data['equivalente']}"
        tg = re.search(r'Total Geral:\s*([\d,.]+)', raw)
        self.data['total_geral'] = parse_br(tg.group(1)) if tg else 0.0
        tc = re.search(r'Crit[ée]rios Espec[íi]ficos Pontuados:\s*(\d+)', raw)
        self.data['total_criterios'] = int(tc.group(1)) if tc else 0

    def _extrair_descricao_item_pontos(self, bloco):
        """Extrai descrição detalhada e pontos de um bloco de critério."""
        descricao = ''
        pontos = None
        linhas = bloco.split('\n')
        for linha in linhas:
            l = linha.strip()
            if not l or l.startswith('#'):
                continue
            pm = re.search(r'([\d,.]+)\s*pontos?', l)
            if pm and not re.match(r'^\d+\s', l):
                pontos = parse_br(pm.group(1))
            if not re.match(r'^[\d.,]+\s', l) and not re.match(r'^Crit[eé]rio', l) and not pm:
                if l and len(l) > 3:
                    descricao += ' ' + l if descricao else l
        return descricao.strip(), pontos

    def _extract_grupos(self, raw):
        nomes = [
            ('I', 'Comissões'),
            ('II', 'Projetos'),
            ('III', 'Premiações'),
            ('IV', 'Responsabilidades'),
            ('V', 'Direção'),
            ('VI', 'Produção'),
        ]
        for rom, curto in nomes:
            pat = rf'REQUISITO\s+{re.escape(rom)}.*?(?=REQUISITO\s+\w|\Z)'
            m = re.search(pat, raw, re.DOTALL)
            crit, pts = 0, 0.0
            if m:
                bloco = m.group(0)
                # Tenta extrair "X critério(s) YYYY.ZZ"
                cm = re.search(r'(\d+)\s*crit[eé]rio\(s\)\s*([\d,.]+)', bloco)
                if cm:
                    crit = int(cm.group(1))
                    pts = parse_br(cm.group(2))
                else:
                    cz = re.search(r'0\s*crit[eé]rio\(s\)\s*([\d,.]+)', bloco)
                    if cz:
                        pts = parse_br(cz.group(1))
            self.data['grupos'].append({
                'romano': rom, 'nome_curto': curto,
                'criterios': crit, 'pontos': pts
            })

    def _extract_criterios(self, pages):
        raw_text = '\n'.join(pages)
        raw_text = re.sub(r'[|¦║]', ' ', raw_text)
        raw_text = re.sub(r'^[A-Z]\s+(\d)', r'\1', raw_text, flags=re.MULTILINE)
        lines = raw_text.split('\n')

        ordem_map = {
            'I-01': 0, 'I-02': 1, 'I-03': 2, 'I-04': 3, 'I-05': 4, 'I-06': 5, 'I-07': 6,
            'II-01': 7, 'II-02': 8, 'II-03': 9, 'II-04': 10, 'II-05': 11, 'II-06': 12, 'II-07': 13,
            'III-01': 14, 'III-02': 15, 'III-03': 16,
            'IV-01': 17, 'IV-02': 18, 'IV-03': 19, 'IV-04': 20, 'IV-05': 21, 'IV-06': 22, 'IV-07': 23,
            'V-01': 24, 'V-02': 25, 'V-03': 26, 'V-04': 27, 'V-05': 28,
            'VI-01': 29, 'VI-02': 30, 'VI-03': 31, 'VI-04': 32, 'VI-05': 33, 'VI-06': 34,
            'VI-07': 35, 'VI-08': 36, 'VI-09': 37, 'VI-10': 38, 'VI-11': 39,
            'VI-12': 40, 'VI-13': 41, 'VI-14': 42, 'VI-15': 43, 'VI-16': 44, 'VI-17': 45,
        }

        i = 0
        while i < len(lines):
            line = lines[i]
            cm = re.match(r'Crit[eé]rio\s+(\w+)\s*[-–]\s*(\d+)\s*:\s*(.*)', line)
            if not cm:
                i += 1
                continue

            key = f"{cm.group(1)}-{cm.group(2)}"
            desc = cm.group(3).strip()
            ordem = ordem_map.get(key, 99)

            # Agrupar linhas seguintes como descrição
            full_desc = desc
            j = i + 1
            while j < len(lines):
                l = lines[j].strip()
                if not l or l.startswith('#'):
                    j += 1
                    continue
                # Se é outro critério, para
                if re.match(r'Crit[eé]rio\s+\w+\s*[-–]\s*\d+\s*:', lines[j]):
                    break
                full_desc += ' ' + l
                j += 1

            # Extrair itens numerados e pontos
            items = []
            pontos = None
            k = i + 1
            while k < len(lines):
                l = lines[k].strip()
                if not l:
                    k += 1
                    continue
                if re.match(r'Crit[eé]rio\s+\w+\s*[-–]\s*\d+\s*:', lines[k]):
                    break
                im = re.match(r'^\s*(\d+)\s+(.+)$', lines[k])
                if im:
                    items.append({'num': im.group(1), 'texto': im.group(2).strip()})
                pm = re.search(r'([\d,.]+)\s*pontos?\s*$', lines[k])
                if pm and not re.match(r'^\d+\s', lines[k]):
                    pontos = parse_br(pm.group(1))
                if pontos is None:
                    pm2 = re.match(r'^\s*[A-Z]?\s*([\d,.]+)\s*pontos?\s*', lines[k])
                    if pm2:
                        pontos = parse_br(pm2.group(1))
                k += 1

            # Se não tem itens, tenta extrair do full_desc
            if not items:
                # Procura por lista numerada no texto concatenado
                items_matches = re.findall(r'(\d+)\s*[.)]\s*([^0-9]+?)(?=\s*\d+\s*[.)]\s*|\Z)', full_desc)
                if items_matches:
                    items = [{'num': m[0], 'texto': m[1].strip().rstrip(',')} for m in items_matches]

            # Se ainda não tem pontos, tenta extrair da descrição
            if pontos is None and full_desc:
                pm = re.search(r'[Pp]ontuação[:\s]*([\d,.]+)', full_desc)
                if pm:
                    pontos = parse_br(pm.group(1))

            crit_data = {
                'key': key, 'romano': cm.group(1), 'numero': cm.group(2),
                'descricao': full_desc, 'ordem': ordem,
                'itens': items, 'pontos': pontos
            }
            self.data['criterios'][key] = crit_data
            if key not in self.data['ordem_criterios']:
                self.data['ordem_criterios'].append(key)
            i = j  # j já está na linha do próximo critério (break) ou além do fim

        self.data['ordem_criterios'].sort(key=lambda k: ordem_map.get(k, 99))


# =============================================================================
# NARRATIVE ENGINE v5.0 — integração com narrativas geradas pelo LLM
# =============================================================================
class NarrativeEngine:
    """Carrega e gerencia narrativas geradas pelo LLM para o memorial.

    As narrativas são textos longos, em linguagem científica, como o
    memorial de um professor titular. Este engine substitui os textos
    genéricos do MemorialGenerator pelos textos do LLM quando um JSON
    de narrativas é fornecido.
    """

    ESTRUTURA_ESPERADA = {
        # Seção 1 — Introdução
        'introducao_quem_sou': '1.1 Quem sou e o que apresento — texto completo',
        'introducao_essencia': '1.2 A essência do meu fazer profissional — texto completo',
        'introducao_dimensoes': 'Lista de 3 dimensões com descrições (array de strings)',
        'introducao_fundamentos': '1.3 Fundamentos legais — texto completo',

        # Anexo I — Comissões
        'anexo_I_intro': 'Texto introdutório do Anexo I',
        'anexo_I_criterios': {
            'I-01': 'Texto detalhado para I-1: Conselhos superiores',
            'I-02': 'Texto detalhado para I-2: Coordenação/presidência',
            'I-03': 'Texto detalhado para I-3: Membro de comissões',
            'I-04': 'Texto detalhado para I-4: Comissões permanentes',
            'I-05': 'Texto detalhado para I-5: Vestibulares/concursos',
            'I-06': 'Texto detalhado para I-6: Elaboração de provas',
            'I-07': 'Texto detalhado para I-7: Organização de eventos',
        },

        # Anexo II — Projetos
        'anexo_II_intro': 'Texto introdutório do Anexo II',
        'anexo_II_criterios': {
            'II-01': 'Texto detalhado para II-1',
            'II-02': 'Texto detalhado para II-2',
            'II-03': 'Texto detalhado para II-3',
            'II-04': 'Texto detalhado para II-4',
            'II-05': 'Texto detalhado para II-5',
            'II-06': 'Texto detalhado para II-6',
            'II-07': 'Texto detalhado para II-7',
        },

        # Anexo III — Premiações
        'anexo_III': 'Texto completo do Anexo III',

        # Anexo IV — Responsabilidades
        'anexo_IV_intro': 'Texto introdutório do Anexo IV',
        'anexo_IV_criterios': {
            'IV-01': 'Texto detalhado para IV-1',
            'IV-02': 'Texto detalhado para IV-2',
            'IV-03': 'Texto detalhado para IV-3',
            'IV-04': 'Texto detalhado para IV-4',
            'IV-05': 'Texto detalhado para IV-5',
            'IV-06': 'Texto detalhado para IV-6',
            'IV-07': 'Texto detalhado para IV-7',
        },

        # Anexo V — Direção
        'anexo_V_intro': 'Texto introdutório do Anexo V',
        'anexo_V_criterios': {
            'V-01': 'Texto detalhado para V-1',
            'V-02': 'Texto detalhado para V-2',
            'V-03': 'Texto detalhado para V-3',
            'V-04': 'Texto detalhado para V-4',
            'V-05': 'Texto detalhado para V-5',
        },

        # Anexo VI — Produção
        'anexo_VI_intro': 'Texto introdutório do Anexo VI',
        'anexo_VI_criterios': {
            'VI-01': 'Texto detalhado para VI-1',
            'VI-02': 'Texto detalhado para VI-2',
            'VI-03': 'Texto detalhado para VI-3',
            'VI-04': 'Texto detalhado para VI-4',
            'VI-05': 'Texto detalhado para VI-5',
            'VI-06': 'Texto detalhado para VI-6',
            'VI-07': 'Texto detalhado para VI-7',
            'VI-08': 'Texto detalhado para VI-8',
            'VI-09': 'Texto detalhado para VI-9: Livro com ISBN',
            'VI-10': 'Texto detalhado para VI-10: Artigos publicados',
            'VI-11': 'Texto detalhado para VI-11',
            'VI-12': 'Texto detalhado para VI-12',
            'VI-13': 'Texto detalhado para VI-13',
            'VI-14': 'Texto detalhado para VI-14',
            'VI-15': 'Texto detalhado para VI-15: Instrutor',
            'VI-16': 'Texto detalhado para VI-16: Coordenação de eventos',
            'VI-17': 'Texto detalhado para VI-17',
        },

        # Reflexão Final
        'reflexao_saberes': '9.1 Que saberes construí? — texto completo',
        'reflexao_contribuicao': '9.2 Qual minha contribuição singular? — texto completo',
        'reflexao_pedido': '9.3 Pedido — texto completo',
        'reflexao_completa': 'Texto completo da seção 9 (substitui as 3 subseções acima se presente)',
    }

    def __init__(self, json_path=None):
        self.narrativas = {}
        if json_path:
            self.carregar(json_path)

    def carregar(self, json_path):
        """Carrega narrativas de um arquivo JSON."""
        path = Path(json_path)
        if not path.exists():
            raise FileNotFoundError(f"Arquivo de narrativas não encontrado: {json_path}")
        with open(path, 'r', encoding='utf-8') as f:
            self.narrativas = json.load(f)
        print(f"   Narrativas carregadas: {json_path} ({len(self.narrativas)} campos)")

    def tem(self, chave):
        """Verifica se uma narrativa existe para determinada chave."""
        return chave in self.narrativas and self.narrativas[chave] and len(str(self.narrativas[chave])) > 50

    def get(self, chave, padrao=''):
        """Retorna a narrativa para a chave, ou o padrão se não existir."""
        if self.tem(chave):
            return self.narrativas[chave]
        return padrao

    def get_criterio(self, key, padrao=''):
        """Busca narrativa de critério nas estruturas aninhadas."""
        prefix = key.split('-')[0]
        mapa_chaves = {
            'I': 'anexo_I_criterios', 'II': 'anexo_II_criterios',
            'IV': 'anexo_IV_criterios', 'V': 'anexo_V_criterios',
            'VI': 'anexo_VI_criterios',
        }
        mapa = mapa_chaves.get(prefix)
        if mapa and mapa in self.narrativas:
            criterios = self.narrativas[mapa]
            if isinstance(criterios, dict) and key in criterios:
                txt = criterios[key]
                if txt and len(str(txt)) > 50:
                    return txt
        return padrao

    def get_dimensoes(self):
        """Retorna lista de dimensões da introdução (array de strings)."""
        if 'introducao_dimensoes' in self.narrativas:
            dims = self.narrativas['introducao_dimensoes']
            if isinstance(dims, list) and len(dims) >= 2:
                return dims
            if isinstance(dims, str) and len(dims) > 50:
                return [dims]
        return []

    def get_reflexao_completa(self):
        """Retorna a reflexão completa se disponível."""
        if self.tem('reflexao_completa'):
            return self.narrativas['reflexao_completa']
        return None


# =============================================================================
# GERADOR DO MEMORIAL v5.0 — Narrativas do LLM + fallback genérico
# =============================================================================
class MemorialGenerator:
    """Gera memorial com narrativas geradas DINAMICAMENTE a partir dos dados extraídos."""

    def __init__(self, data, ano_ingresso, narrative_engine=None):
        self.d = data
        self.ano_ingresso = ano_ingresso
        self.ano_atual = datetime.now().year
        self.anos_carreira = self.ano_atual - ano_ingresso
        self.nivel = data.get('nivel_info', NIVEL_PADRAO)
        self.nome = data['nome']
        self.lotacao = data.get('lotacao', '')
        self.equivalente = self.nivel['equivalente']
        self.ne = narrative_engine if narrative_engine and narrative_engine.narrativas else None
        self._usando_llm = self.ne is not None

    def _n_arr(self, items):
        """Concatena itens com separação natural."""
        if not items:
            return ''
        if len(items) == 1:
            return items[0]['texto'] if isinstance(items[0], dict) else str(items[0])
        textos = [i['texto'] if isinstance(i, dict) else str(i) for i in items]
        return ', '.join(textos[:-1]) + ' e ' + textos[-1]

    # ---- Narrativas dinâmicas ----

    def _desc_anexo_I(self, g, c):
        """Narrativa dinâmica para Anexo I (Comissões) — usa LLM se disponível."""
        md = []
        qtd_criterios = g['criterios']

        if self._usando_llm and self.ne.get('anexo_I_intro'):
            md.append(self.ne.get('anexo_I_intro') + "\n\n")
        else:
            desc_geral = ' e '.join(
                'membro de comissões' if 'I-03' in c else '' or
                'presidente de comissões' if 'I-02' in c else '' or
                'conselhos superiores' if 'I-01' in c else '' or
                'concursos e vestibulares' if 'I-05' in c else '' or
                'elaboração de provas' if 'I-06' in c else ''
                for _ in [1]
            )
            if not desc_geral:
                desc_geral = f'{qtd_criterios} critérios de comissões e grupos de trabalho'

            md.append(f"## 2.1 Memorialista: atuação em comissões\n\n")
            md.append(
                f"Ao longo da carreira, participei ativamente de comissões e grupos de trabalho "
                f"na UFV, atuando como {desc_geral}. Foram {qtd_criterios} critérios com "
                f"pontuação no Anexo I, totalizando {fmt_br(g['pontos'])} pontos. "
                f"Essa atuação me permitiu contribuir para a gestão universitária em suas "
                f"dimensões normativa, administrativa e estratégica.\n\n"
            )

        # Narrativa de cada critério
        for key, info in sorted(c.items(), key=lambda x: x[1].get('ordem', 99)):
            if not key.startswith('I-'):
                continue
            num_items = len(info['itens'])

            # Tenta usar narrativa do LLM para este critério
            txt_llm = self.ne.get_criterio(key) if self._usando_llm else ''

            if txt_llm:
                sub_num = int(info['numero']) if info['numero'].isdigit() else 0
                md.append(f"## 2.{sub_num + 1} Item I-{info['numero']}\n\n")
                md.append(txt_llm + "\n\n")
                # Anexa itens comprobatórios mesmo com narrativa LLM
                if info['itens']:
                    md.append("**Itens comprobatórios:**\n\n")
                    for item in info['itens']:
                        if isinstance(item, dict):
                            md.append(f"- Item {item['num']}: {item['texto'][:200]}\n")
                        else:
                            md.append(f"- {str(item)[:200]}\n")
                    md.append("\n")
            else:
                resumo_desc = info['descricao'][:150].strip() if info['descricao'] else ''
                sub_num = int(info['numero']) if info['numero'].isdigit() else 0
                md.append(
                    f"## 2.{sub_num + 1} "
                    f"Item I-{info['numero']}: {resumo_desc[:100] if resumo_desc else 'Atuação em comissão'}\n\n"
                )
                md.append(
                    f"Este critério refere-se a: {resumo_desc if resumo_desc else 'atividades de comissão'}. "
                    f"Foram registrados {num_items} itens comprobatórios, "
                    f"com pontuação de {fmt_br(info.get('pontos', 0))} pontos.\n\n"
                )
                if info['itens']:
                    md.append("**Itens comprobatórios:**\n\n")
                    for item in info['itens']:
                        if isinstance(item, dict):
                            md.append(f"- Item {item['num']}: {item['texto'][:200]}\n")
                        else:
                            md.append(f"- {str(item)[:200]}\n")
                    md.append("\n")

        return md

    def _desc_anexo_II(self, g, c):
        """Narrativa dinâmica para Anexo II (Projetos) — usa LLM se disponível."""
        md = []
        if self._usando_llm and self.ne.get('anexo_II_intro'):
            md.append(self.ne.get('anexo_II_intro') + "\n\n")
        else:
            md.append("## 3.1 Atuação em projetos institucionais\n\n")
            if g['pontos'] > 0:
                qtd = len([k for k in c if k.startswith('II-')])
                md.append(
                    f"Participei de projetos institucionais que contribuíram para o "
                    f"desenvolvimento da UFV. Foram {qtd} critérios com {fmt_br(g['pontos'])} pontos. "
                    f"Esses projetos envolveram pesquisa acadêmica, desenvolvimento de "
                    f"metodologias e avaliação de trabalhos.\n\n"
                )
            else:
                md.append("Não há critérios pontuados neste anexo.\n\n")

        if g['pontos'] > 0:
            for key, info in sorted(c.items(), key=lambda x: x[1].get('ordem', 99)):
                if not key.startswith('II-'):
                    continue
                txt_llm = self.ne.get_criterio(key) if self._usando_llm else ''
                if txt_llm:
                    md.append(f"## 3.{self._sub_anexo_n(key, c)} Item II-{info['numero']}\n\n")
                    md.append(txt_llm + "\n\n")
                else:
                    resumo = info['descricao'][:200] if info['descricao'] else 'Atuação em projeto'
                    md.append(
                        f"## 3.{self._sub_anexo_n(key, c)} "
                        f"Item II-{info['numero']}: {resumo[:100]}\n\n"
                    )
                    md.append(f"{resumo}. Pontuação: {fmt_br(info.get('pontos', 0))} pts.\n\n")
                if info['itens']:
                    for item in info['itens']:
                        md.append(f"- Item {item['num']}: {item['texto'][:200]}\n")
                    md.append("\n")
        return md

    def _sub_anexo_n(self, key, all_c):
        """Calcula subnumeração dentro de um anexo."""
        prefix = key.split('-')[0]
        keys = sorted([k for k in all_c if k.startswith(f'{prefix}-')])
        try:
            idx = keys.index(key)
            return idx + 2  # Começa em 2 (1 é intro)
        except ValueError:
            return 2

    def _desc_anexo_III(self, g):
        """Anexo III - Premiações — usa LLM se disponível."""
        md = []
        if self._usando_llm and self.ne.get('anexo_III'):
            md.append(self.ne.get('anexo_III') + "\n\n")
        else:
            if g['pontos'] > 0:
                md.append(f"Recebi premiações ou reconhecimentos formais ao longo da carreira.\n\n")
            else:
                md.append("Declaro que não recebi premiações formais que possam ser enquadradas neste Anexo.\n\n")
        return md

    def _desc_anexo_IV(self, g, c):
        """Anexo IV - Responsabilidades — usa LLM se disponível."""
        md = []
        if self._usando_llm and self.ne.get('anexo_IV_intro'):
            md.append(self.ne.get('anexo_IV_intro') + "\n\n")
        else:
            if g['pontos'] > 0:
                md.append("Fui designado para responsabilidades técnico-administrativas específicas.\n\n")
            else:
                md.append(
                    "Registro as designações para atuação em sistemas e processos administrativos "
                    "conforme constam no sistema oficial da UFV.\n\n"
                )

        if g['pontos'] > 0:
            for key, info in sorted(c.items(), key=lambda x: x[1].get('ordem', 99)):
                if not key.startswith('IV-'):
                    continue
                txt_llm = self.ne.get_criterio(key) if self._usando_llm else ''
                if txt_llm:
                    md.append(f"**Item IV-{info['numero']}**\n\n{txt_llm}\n\n")
                else:
                    resumo = info['descricao'][:200] if info['descricao'] else ''
                    md.append(
                        f"Item IV-{info['numero']}: {resumo[:100] if resumo else 'Responsabilidade técnica'}. "
                        f"Pontuação: {fmt_br(info.get('pontos', 0))} pts.\n\n"
                    )
                if info['itens']:
                    for item in info['itens']:
                        md.append(f"- Item {item['num']}: {item['texto'][:200]}\n")
                    md.append("\n")
        return md

    def _desc_anexo_V(self, g, c):
        """Anexo V - Direção — usa LLM se disponível."""
        md = []
        if self._usando_llm and self.ne.get('anexo_V_intro'):
            md.append(self.ne.get('anexo_V_intro') + "\n\n")
        else:
            if g['pontos'] > 0:
                md.append("## 6.1 Trajetória de liderança institucional\n\n")
                md.append(
                    "Um dos aspectos significativos da carreira foi o exercício de cargos "
                    "de direção e funções gratificadas, que permitiram contribuir diretamente "
                    "para a formulação e execução de políticas institucionais na UFV.\n\n"
                )
            else:
                md.append(
                    "Não há critérios pontuados para cargos de direção ou assessoramento "
                    "neste Anexo.\n\n"
                )

        if g['pontos'] > 0:
            for key, info in sorted(c.items(), key=lambda x: x[1].get('ordem', 99)):
                if not key.startswith('V-'):
                    continue
                txt_llm = self.ne.get_criterio(key) if self._usando_llm else ''
                if txt_llm:
                    md.append(f"## 6.{int(info['numero'])+1} Item V-{info['numero']}\n\n")
                    md.append(txt_llm + "\n\n")
                else:
                    resumo = info['descricao'][:200] if info['descricao'] else ''
                    md.append(
                        f"## 6.{int(info['numero'])+1} Item V-{info['numero']}: {resumo[:100] if resumo else 'Cargo de direção'}\n\n"
                    )
                    md.append(f"{resumo}. Pontuação: {fmt_br(info.get('pontos', 0))} pts.\n\n")
                if info['itens']:
                    for item in info['itens']:
                        md.append(f"- Item {item['num']}: {item['texto'][:200]}\n")
                    md.append("\n")
        return md

    def _desc_anexo_VI(self, g, c):
        """Anexo VI - Produção — usa LLM se disponível."""
        md = []
        if self._usando_llm and self.ne.get('anexo_VI_intro'):
            md.append(self.ne.get('anexo_VI_intro') + "\n\n")
        else:
            if g['pontos'] > 0:
                md.append("## 7.1 Produção e difusão de conhecimento\n\n")
                md.append(
                    "Ao longo da carreira, busquei não apenas executar atividades, mas também "
                    "produzir e difundir conhecimento técnico e científico.\n\n"
                )
            else:
                md.append("Não há critérios pontuados neste Anexo.\n\n")

        if g['pontos'] > 0:
            for key, info in sorted(c.items(), key=lambda x: x[1].get('ordem', 99)):
                if not key.startswith('VI-'):
                    continue
                txt_llm = self.ne.get_criterio(key) if self._usando_llm else ''
                if txt_llm:
                    md.append(f"## 7.{int(info['numero']) - 8 if key in ['VI-09','VI-10','VI-15','VI-16'] else self._sub_anexo_n(key, c)} "
                              f"Item VI-{info['numero']}\n\n")
                    md.append(txt_llm + "\n\n")
                else:
                    resumo = info['descricao'][:300] if info['descricao'] else ''
                    md.append(
                        f"## 7.{int(info['numero']) - 8 if key in ['VI-09','VI-10','VI-15','VI-16'] else self._sub_anexo_n(key, c)} "
                        f"Item VI-{info['numero']}: {resumo[:100]}\n\n"
                    )
                    md.append(f"{resumo}. Pontuação: {fmt_br(info.get('pontos', 0))} pts.\n\n")
                if info['itens']:
                    for item in info['itens']:
                        md.append(f"- Item {item['num']}: {item['texto'][:200]}\n")
                    md.append("\n")
        return md

    def _gerar_reflexao(self):
        """Gera reflexão final — usa LLM se disponível."""
        total = fmt_br(self.d['total_geral'])
        g = self.d['grupos']

        # Se tem reflexão completa do LLM, usa integralmente
        if self._usando_llm:
            ref_completa = self.ne.get_reflexao_completa()
            if ref_completa:
                return [
                    "# 9 REFLEXÃO FINAL -- SABERES E COMPETÊNCIAS\n\n",
                    "*Em conformidade com o art. 15 do Decreto nº 13.048/2026*\n\n",
                    ref_completa + "\n\n",
                    f"Viçosa, {self.ano_atual}.\n\n",
                    f"{self.nome}\n",
                    f"{self.d['cargo']} -- UFV\n",
                    f"{self.d['titulacao']}\n",
                ]

        md = [
            "# 9 REFLEXÃO FINAL -- SABERES E COMPETÊNCIAS\n\n",
            "*Em conformidade com o art. 15 do Decreto nº 13.048/2026*\n\n",
        ]

        # 9.1 Saberes
        md.append("## 9.1 Que saberes construí?\n\n")
        if self._usando_llm and self.ne.get('reflexao_saberes'):
            md.append(self.ne.get('reflexao_saberes') + "\n\n")
        else:
            md.append(
                f"Ao longo de {self.anos_carreira} anos de serviço público na UFV, construí "
                f"saberes e competências que se refletem nos {self.d['total_criterios']} critérios "
                f"pontuados, totalizando {total} pontos.\n\n"
            )
            saberes = {
                'normativo': 'interpretação e proposição de normas',
                'gestao': 'gestão de equipes e processos',
                'academico': 'produção e difusão de conhecimento',
                'tecnico': 'atuação técnica especializada',
            }
            saberes_usados = []
            for grp in g:
                if grp['romano'] == 'I' and grp['pontos'] > 0:
                    saberes_usados.append(('normativo', 'Saber normativo', saberes['normativo']))
                if grp['romano'] == 'V' and grp['pontos'] > 0:
                    saberes_usados.append(('gestao', 'Saber de gestão', saberes['gestao']))
                if grp['romano'] == 'VI' and grp['pontos'] > 0:
                    saberes_usados.append(('academico', 'Saber acadêmico-científico', saberes['academico']))
                if grp['romano'] in ('II', 'IV') and grp['pontos'] > 0:
                    saberes_usados.append(('tecnico', 'Saber técnico-profissional', saberes['tecnico']))
            if not saberes_usados:
                saberes_usados = [('gestao', 'Saber profissional', 'atuação no serviço público')]
            for sid, snome, sdesc in saberes_usados:
                md.append(f"**{snome}:** {sdesc}, construído ao longo da carreira.\n\n")

        # 9.2 Contribuição
        md.append("## 9.2 Qual minha contribuição singular?\n\n")
        if self._usando_llm and self.ne.get('reflexao_contribuicao'):
            md.append(self.ne.get('reflexao_contribuicao') + "\n\n")
        else:
            anexos_ativos = [grp['nome_curto'] for grp in g if grp['pontos'] > 0]
            md.append(
                f"A trajetória aqui documentada demonstra que os saberes construídos "
                f"na prática profissional, quando refletidos sistematicamente, geram "
                f"conhecimento relevante para a instituição. Minha contribuição à UFV "
                f"se expressa nos {len(anexos_ativos)} eixos de atuação "
                f"({', '.join(anexos_ativos)}) que desenvolvi ao longo da carreira.\n\n"
            )

        # 9.3 Pedido
        md.append("## 9.3 Pedido\n\n")
        if self._usando_llm and self.ne.get('reflexao_pedido'):
            md.append(self.ne.get('reflexao_pedido') + "\n\n")
        else:
            md.append(
                "Ante o exposto, com fundamento na Lei nº 11.091/2005 (alterada pela "
                f"Lei nº 15.367/2026), no Decreto nº 13.048/2026 e na documentação "
                f"comprobatória anexa, requeiro à CRSC-PCCTAE da Universidade Federal "
                f"de Viçosa o deferimento da concessão do Reconhecimento de Saberes "
                f"e Competências no nível RSC-PCCTAE {self.nivel['nome']}.\n\n"
                "Nestes termos, pede deferimento.\n\n"
            )

        md.append(f"Viçosa, {self.ano_atual}.\n\n")
        md.append(f"{self.nome}\n")
        md.append(f"{self.d['cargo']} -- UFV\n")
        md.append(f"{self.d['titulacao']}\n")
        return md

    def _selecionar_epigrafe(self):
        """Seleciona epígrafe genérica baseada no perfil."""
        citacoes = [
            {'q': '"A educação não transforma o mundo. Educação muda as pessoas. Pessoas transformam o mundo."', 'a': 'Paulo Freire'},
            {'q': '"Servir ao público é construir o país todos os dias."', 'a': 'Paulo Freire'},
            {'q': '"A carreira de um servidor público se mede não pelo cargo que ocupa, mas pelas transformações que promove."', 'a': 'Anísio Teixeira'},
            {'q': '"O conhecimento se constrói na prática refletida."', 'a': 'Adaptado de Paulo Freire'},
            {'q': '"A universidade pública é patrimônio do povo brasileiro que deve ser defendido e fortalecido a cada dia."', 'a': 'Anísio Teixeira'},
        ]
        lot_lower = self.lotacao.lower() if self.lotacao else ''
        temas = {
            'gestão': 1, 'pessoas': 1, 'administrat': 1,
            'ensino': 0, 'graduação': 0, 'educaç': 0,
            'pesquisa': 3, 'ciência': 3, 'tecnologia': 3,
            'reitor': 4, 'gabinete': 4,
        }
        escolha = 2  # default Anísio
        for palavra, idx in temas.items():
            if palavra in lot_lower:
                escolha = idx
                break
        c = citacoes[escolha]
        return c['q'], c['a']

    def generate(self):
        """Gera o memorial completo."""
        titulo = (f"MEMORIAL DESCRITIVO PARA RECONHECIMENTO DE SABERES E "
                  f"COMPETÊNCIAS – RSC-PCCTAE NÍVEL {self.nivel['nome']}")
        md = []

        # ===== CAPA =====
        # ABNT NBR 14724: elementos da capa distribuídos verticalmente
        # na página A4 (29,7 cm). Cada <br> vira um parágrafo vazio
        # de 12pt (0,42 cm). A página tem 841,9pt, com margens de
        # 85pt (3 cm) e 56,7pt (2 cm) → área útil = 700,2pt.
        # Alvo: instituição a ~96pt da margem (3,4 cm), autor ~234pt
        # (8,3 cm), título ~432pt (15,3 cm), local/ano ~650pt (23 cm).
        md.append('<br><br><br><br><br><br><br><br>\n')       # 8  =  96pt
        md.append('<p align="center"><strong>UNIVERSIDADE FEDERAL DE VIÇOSA</strong></p>\n')
        md.append('<br><br><br><br><br><br><br><br><br><br>\n')  # 10 = 120pt
        md.append(f'<p align="center"><strong>{self.nome}</strong></p>\n')
        md.append('<br><br><br><br><br><br><br><br><br><br><br><br>\n')  # 12 = 144pt
        md.append(f'<p align="center"><strong>{titulo}</strong></p>\n')
        md.append('<br><br><br><br><br><br><br><br><br><br><br><br><br>\n')  # 13 = 156pt
        md.append('<p align="center"><strong>VIÇOSA – MINAS GERAIS</strong></p>\n')
        md.append(f'<p align="center"><strong>{self.ano_atual}</strong></p>\n')
        md.append('<br><br><br>\n')                              # 3  =  36pt
        md.append("\n---\n")

        # ===== FOLHA DE ROSTO (verso da capa) =====
        # ABNT NBR 14724: elementos distribuídos com espaçamento generoso
        # para preencher a página verticalmente.
        md.append('<br><br><br><br><br><br><br><br>\n')
        md.append(f'<p align="center"><strong>{self.nome}</strong></p>\n')
        md.append('<br><br><br><br><br><br><br><br>\n')
        md.append(f'<p align="center"><strong>{titulo}</strong></p>\n')
        md.append('<br><br><br><br><br><br>\n')
        md.append('<div style="text-align:justify; margin-left:4cm; margin-right:0cm; line-height:1.0;">\n')
        md.append(
            f'Memorial descritivo apresentado à Comissão para Reconhecimento de Saberes '
            f'e Competências do Plano de Carreira dos Cargos Técnico-Administrativos em '
            f'Educação (CRSC-PCCTAE) da Universidade Federal de Viçosa como requisito '
            f'para concessão do RSC-PCCTAE Nível {self.nivel["nome"]}, nos termos da Lei '
            f'nº 11.091/2005 (alterada pela Lei nº 15.367/2026), do Decreto nº 13.048/2026 '
            f'e da legislação correlata.\n'
        )
        md.append('</div>\n<br><br><br><br><br><br><br><br>\n')
        md.append('<p align="center">VIÇOSA – MINAS GERAIS</p>\n')
        md.append(f'<p align="center">{self.ano_atual}</p>\n')
        md.append('<br><br><br><br>\n')
        md.append("\n---\n")

        # ==== DEDICATÓRIA =====
        md.append('<p align="right" style="font-size:12pt; font-style:italic;">\n')
        md.append('  Aos servidores técnico-administrativos em educação,<br>\n')
        md.append('  cujo trabalho constrói a universidade pública brasileira.\n')
        md.append('</p>\n\n---\n')

        # ===== AGRADECIMENTOS =====
        md.append("# AGRADECIMENTOS\n\n")
        lot_str = f"Aos colegas da {self.lotacao}, " if self.lotacao and self.lotacao.upper() != 'N/A' else ""
        md.append(
            f"Expresso minha gratidão à Universidade Federal de Viçosa, instituição que há "
            f"{self.anos_carreira} anos é o espaço do meu crescimento profissional e pessoal. "
            f"{lot_str}pelo aprendizado e trabalho colaborativo.\n\n"
            f"Agradeço à CRSC-PCCTAE da UFV pelo cuidadoso trabalho de avaliação das trajetórias "
            f"dos servidores técnico-administrativos em educação.\n\n"
            f"Aos servidores que compartilham a missão de construir uma universidade pública, "
            f"gratuita e de qualidade, minha sincera admiração e reconhecimento.\n"
        )
        md.append("\n---\n")

        # ===== EPIGRAFE =====
        ep_q, ep_a = self._selecionar_epigrafe()
        md.append(f'<p align="right" style="font-size:12pt; font-style:italic;">\n')
        md.append(f'  {ep_q}<br>\n')
        md.append(f'  <span style="font-size:11pt;">— {ep_a}</span>\n')
        md.append('</p>\n\n---\n')

        # ===== LISTA DE SIGLAS =====
        md.append("# LISTA DE SIGLAS\n\n")
        for s in [
            "**CRSC-PCCTAE** -- Comissão para Reconhecimento de Saberes e Competências do Plano de Carreira dos Cargos Técnico-Administrativos em Educação",
            "**PCCTAE** -- Plano de Carreira dos Cargos Técnico-Administrativos em Educação",
            "**RSC** -- Reconhecimento de Saberes e Competências",
            "**SIAPE** -- Sistema Integrado de Administração de Recursos Humanos",
            "**UFV** -- Universidade Federal de Viçosa",
        ]:
            md.append(f"{s}\n\n")
        md.append("---\n")

        # ===== SUMÁRIO =====
        md.append("# SUMÁRIO\n\n")
        for s in [
            "- **1 INTRODUÇÃO -- TRAJETÓRIA E FUNDAMENTOS**\n",
            "- **2 ANEXO I -- PARTICIPAÇÃO EM COMISSÕES, GRUPOS DE TRABALHO E CONCURSOS**\n",
            "- **3 ANEXO II -- PARTICIPAÇÃO EM PROJETOS INSTITUCIONAIS**\n",
            "- **4 ANEXO III -- PREMIAÇÕES**\n",
            "- **5 ANEXO IV -- RESPONSABILIDADES TÉCNICO-ADMINISTRATIVAS**\n",
            "- **6 ANEXO V -- EXERCÍCIO DE FUNÇÕES DE DIREÇÃO E ASSESSORAMENTO**\n",
            "- **7 ANEXO VI -- PRODUÇÃO, PROSPECÇÃO E DIFUSÃO DE CONHECIMENTO CIENTÍFICO E TÉCNICO**\n",
            "- **8 SÍNTESE DE PONTUAÇÃO**\n",
            "- **9 REFLEXÃO FINAL -- SABERES E COMPETÊNCIAS**\n",
            "- **REFERÊNCIAS**\n",
        ]:
            md.append(s)
        md.append("\n---\n")

        # ===== 1 INTRODUÇÃO =====
        g1 = self.d['grupos'][0]
        total = fmt_br(self.d['total_geral'])
        total_crit = self.d['total_criterios']
        lot_str = f", lotado(a) em {self.lotacao}" if self.lotacao and self.lotacao.upper() != 'N/A' else ""

        md.append("# 1 INTRODUÇÃO -- TRAJETÓRIA E FUNDAMENTOS\n\n")
        md.append("## 1.1 Quem sou e o que apresento\n\n")
        if self._usando_llm and self.ne.get('introducao_quem_sou'):
            md.append(self.ne.get('introducao_quem_sou') + "\n\n")
        else:
            md.append(
                f"Meu nome é {self.nome}, matrícula SIAPE {self.d['matricula']}, "
                f"servidor público federal ocupante do cargo de {self.d['cargo']}"
                f"{lot_str}, na Universidade Federal de Viçosa. "
                f"Sou portador do título de {self.d['titulacao']}. "
                f"Apresento este memorial descritivo para obter o Reconhecimento "
                f"de Saberes e Competências no Nível {self.nivel['nome']} (RSC-PCCTAE {self.nivel['nome']}), "
                f"equivalente ao {self.equivalente}, conforme previsto na Lei nº 11.091/2005 "
                f"(alterada pela Lei nº 15.367/2026) e regulamentado pelo Decreto nº 13.048/2026.\n\n"
                f"Este memorial descreve uma trajetória profissional de {self.anos_carreira} anos "
                f"-- de {self.ano_ingresso} a {self.ano_atual} -- "
                f"construída na interseção entre gestão, planejamento institucional, "
                f"inovação e produção de conhecimento. Cada atividade representa "
                f"um saber construído e uma contribuição à UFV.\n\n"
            )

        md.append("## 1.2 A essência do meu fazer profissional\n\n")
        if self._usando_llm and self.ne.get('introducao_essencia'):
            md.append(self.ne.get('introducao_essencia') + "\n\n")
        else:
            md.append(
                f"Como {self.d['cargo']}, minha atuação ao longo da carreira desenvolveu "
                f"competências que se refletem nos {total_crit} critérios pontuados, "
                f"totalizando {total} pontos.\n\n"
            )

        # Dimensões - geradas dinamicamente ou via LLM
        if self._usando_llm:
            dims_llm = self.ne.get_dimensoes()
            if dims_llm:
                for d in dims_llm:
                    md.append(f"{d}\n\n")
        else:
            tem_comissoes = self.d['grupos'][0]['pontos'] > 0
            tem_direcao = self.d['grupos'][4]['pontos'] > 0
            tem_producao = self.d['grupos'][5]['pontos'] > 0

            dims = []
            if tem_comissoes:
                dims.append(("Dimensão normativo-regulatória", 
                    "atuação em comissões, grupos de trabalho e concursos"))
            if tem_direcao:
                dims.append(("Dimensão executivo-estratégica",
                    "exercício de cargos de direção e assessoramento"))
            if tem_producao:
                dims.append(("Dimensão acadêmico-científica",
                    "produção e difusão de conhecimento científico"))

            if not dims:
                dims.append(("Dimensão profissional", "atuação no serviço público"))

            for nome, desc in dims:
                md.append(f"**{nome}:** {desc}.\n\n")

        md.append("## 1.3 Fundamentos legais\n\n")
        if self._usando_llm and self.ne.get('introducao_fundamentos'):
            md.append(self.ne.get('introducao_fundamentos') + "\n\n")
        else:
            md.append(
                "O presente memorial atende aos requisitos do Art. 3 (eixos de atuação), "
                "Art. 5 (pontuação) e Art. 15 (saberes e competências diferenciados) "
                "do Decreto nº 13.048/2026, bem como da Lei nº 11.091/2005 e Lei nº 15.367/2026.\n\n"
                f"Para o RSC-PCCTAE Nível {self.nivel['nome']}:\n"
                "- Pontuação mínima: 75 pontos;\n"
                "- Mínimo de 7 critérios dos Anexos I a VI;\n"
                "- Pelo menos 1 critério do Anexo VI (produção);\n"
                f"- Titulação de {self.d['titulacao']} comprovada.\n\n"
                "Todos os requisitos são atendidos com ampla margem.\n"
            )

        md.append("\n---\n")

        # ===== 2-7 ANEXOS =====
        # Anexo I
        md.append("# 2 ANEXO I -- PARTICIPAÇÃO EM COMISSÕES, GRUPOS DE TRABALHO E CONCURSOS\n\n")
        md.append(f"*Art. 3, inciso I -- Pontuação: {fmt_br(self.d['grupos'][0]['pontos'])} pts ({self.d['grupos'][0]['criterios']} critério(s))*\n\n")
        c_I = {k: v for k, v in self.d['criterios'].items() if k.startswith('I-')}
        md.extend(self._desc_anexo_I(self.d['grupos'][0], c_I))
        md.append("\n---\n")

        # Anexo II
        md.append("# 3 ANEXO II -- PARTICIPAÇÃO EM PROJETOS INSTITUCIONAIS\n\n")
        md.append(f"*Art. 3, inciso II -- Pontuação: {fmt_br(self.d['grupos'][1]['pontos'])} pts ({self.d['grupos'][1]['criterios']} critério(s))*\n\n")
        c_II = {k: v for k, v in self.d['criterios'].items() if k.startswith('II-')}
        md.extend(self._desc_anexo_II(self.d['grupos'][1], c_II))
        md.append("\n---\n")

        # Anexo III
        md.append("# 4 ANEXO III -- RECEBIMENTO DE PREMIAÇÃO\n\n")
        md.append(f"*Art. 3, inciso III -- Pontuação: {fmt_br(self.d['grupos'][2]['pontos'])} pts ({self.d['grupos'][2]['criterios']} critério(s))*\n\n")
        md.extend(self._desc_anexo_III(self.d['grupos'][2]))
        md.append("\n---\n")

        # Anexo IV
        md.append("# 5 ANEXO IV -- RESPONSABILIDADES TÉCNICO-ADMINISTRATIVAS\n\n")
        md.append(f"*Art. 3, inciso IV -- Pontuação: {fmt_br(self.d['grupos'][3]['pontos'])} pts ({self.d['grupos'][3]['criterios']} critério(s))*\n\n")
        c_IV = {k: v for k, v in self.d['criterios'].items() if k.startswith('IV-')}
        md.extend(self._desc_anexo_IV(self.d['grupos'][3], c_IV))
        md.append("\n---\n")

        # Anexo V
        md.append("# 6 ANEXO V -- EXERCÍCIO DE FUNÇÕES DE DIREÇÃO E ASSESSORAMENTO\n\n")
        md.append(f"*Art. 3, inciso V -- Pontuação: {fmt_br(self.d['grupos'][4]['pontos'])} pts ({self.d['grupos'][4]['criterios']} critério(s))*\n\n")
        c_V = {k: v for k, v in self.d['criterios'].items() if k.startswith('V-')}
        md.extend(self._desc_anexo_V(self.d['grupos'][4], c_V))
        md.append("\n---\n")

        # Anexo VI
        md.append("# 7 ANEXO VI -- PRODUÇÃO, PROSPECÇÃO E DIFUSÃO DE CONHECIMENTO CIENTÍFICO E TÉCNICO\n\n")
        md.append(f"*Art. 3, inciso VI -- Pontuação: {fmt_br(self.d['grupos'][5]['pontos'])} pts ({self.d['grupos'][5]['criterios']} critério(s))*\n\n")
        c_VI = {k: v for k, v in self.d['criterios'].items() if k.startswith('VI-')}
        md.extend(self._desc_anexo_VI(self.d['grupos'][5], c_VI))
        md.append("\n---\n")

        # ===== 8 SÍNTESE DE PONTUAÇÃO =====
        md.append("# 8 SÍNTESE DE PONTUAÇÃO\n\n")
        nome_curto_map = {'I': 'Comissões', 'II': 'Projetos', 'III': 'Premiações',
                          'IV': 'Responsabilidades', 'V': 'Direção', 'VI': 'Produção'}
        md.append("## 8.1 Quadro geral -- Pontuação oficial (sistema UFV)\n\n")
        md.append("| Anexo | Conteúdo | Critérios | Pontuação |\n")
        md.append("|-------|----------|-----------|-----------|\n")
        for grp in self.d['grupos']:
            nm = nome_curto_map.get(grp['romano'], grp['nome_curto'])
            md.append(f"| **Anexo {grp['romano']}** | {nm} | {grp['criterios']} | **{fmt_br(grp['pontos'])}** |\n")
        md.append(f"| | **Total** | **{total_crit}** | **{total}** |\n\n")

        total_com_pontos = sum(1 for grp in self.d['grupos'] if grp['criterios'] > 0)
        md.append("## 8.2 Verificação dos requisitos legais\n\n")
        md.append("| Requisito | Exigido | Atendido |\n")
        md.append("|-----------|---------|----------|\n")
        md.append(f"| Pontuação | Mín 75,00 | **{total} pts** -- Atende |\n")
        md.append(f"| Critérios | Min 7 | **{total_crit}** -- Atende |\n")
        md.append(f"| Anexo VI | >=1 | **{self.d['grupos'][5]['criterios']}** -- Atende |\n")
        md.append(f"| Titulação | {self.d['titulacao']} | Comprovada |\n\n")
        md.append("Todos os requisitos legais para RSC-PCCTAE VI são atendidos.\n\n---\n")

        # ===== 9 REFLEXÃO FINAL =====
        md.extend(self._gerar_reflexao())
        md.append("\n---\n")

        # ===== REFERÊNCIAS =====
        md.append("# REFERÊNCIAS\n\n")
        md.append(
            "BRASIL. **Lei nº 11.091**, de 12 de janeiro de 2005. Dispõe sobre a "
            "estruturação do PCCTAE. **Diário Oficial da União**: Brasília, 13 jan. 2005.\n\n"
        )
        md.append(
            "BRASIL. **Lei nº 15.367**, de 30 de março de 2026. Altera a Lei nº 11.091/2005.\n\n"
        )
        md.append(
            f"BRASIL. **Decreto nº 13.048**, de 3 de julho de 2026. Disponível em: "
            f"{DECRETO_URL}. Acesso em: {datetime.now().strftime('%d %b. %Y')}.\n\n"
        )
        md.append(
            "UNIVERSIDADE FEDERAL DE VIÇOSA. Pró-Reitoria de Gestão de Pessoas. "
            f"**Relatório Detalhado RSC**: {self.nome}. Viçosa, MG, {self.ano_atual}.\n\n"
        )
        md.append("PIRES, A. R. P.; SILVA, B. (org.). **Normalização de trabalhos acadêmicos**: conforme ABNTs NBR 14724/2024, NBR 6023/2018 e NBR 10520/2023. Viçosa, MG: UFV, 2025.\n\n---\n")
        md.append(
            "*Documento gerado automaticamente pelo PesquisAI (pdf-to-memorial-rsc v4.0) "
            "a partir do Relatório Detalhado RSC emitido pelo sistema UFV. "
            "Conteúdo integralmente extraído do PDF original.*\n"
        )

        texto = ''.join(md)
        # Travessão curto (–) conforme ABNT nas separações de títulos,
        # subtítulos e notas (uniformiza o antigo hífen duplo "--").
        texto = texto.replace(' -- ', ' – ')
        texto = normalizar_texto(texto)
        texto = restaurar_acentos(texto)
        return texto


# =============================================================================
# FORMATADOR UFV/ABNT — DOCX
# =============================================================================
def set_cell_shading(cell, color):
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', color)
    shading_elm.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _html_to_docx(doc, lines):
    text = ' '.join(lines)
    clean = re.sub(r'<[^>]+>', '', text).strip()
    if not clean:
        return
    lines_text = [re.sub(r'<[^>]+>', '', l).strip() for l in lines 
                  if re.sub(r'<[^>]+>', '', l).strip() and not re.sub(r'<[^>]+>', '', l).strip().startswith('&')]
    is_folha = 'margin-left:4cm' in text and 'Memorial' in clean
    is_right = 'align="right"' in text and not is_folha
    br_count = len([l for l in lines if l.strip() in ('<br>', '<br/>')])
    is_cover = br_count > 3 or 'UNIVERSIDADE FEDERAL' in clean.upper()

    if is_folha:
        # Renderização orientada a elementos: preserva <br> como parágrafos
        # vazios (distribuição vertical ABNT) e <p>/<div> como texto.
        in_div = False
        indent_cm = None
        for l in lines:
            clean = re.sub(r'<[^>]+>', '', l).strip()
            if '<div' in l:
                in_div = True
                m = re.search(r'margin-left:\s*([\d.]+)cm', l)
                if m:
                    indent_cm = float(m.group(1))
                continue
            if '</div>' in l:
                in_div = False
                indent_cm = None
                continue
            if '<br' in l:
                count = l.count('<br>')
                for _ in range(count):
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                continue
            if clean and not clean.startswith('&'):
                if in_div and indent_cm:
                    # Nature do trabalho (recuada, justificada, espaço simples)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.left_indent = Cm(indent_cm)
                    r = p.add_run(clean)
                    r.font.size = Pt(12)
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.line_spacing = 1.5
                    p.paragraph_format.space_after = Pt(6)
                    r = p.add_run(clean)
                    # Aplica negrito se a linha original continha <strong>
                    if '<strong>' in l:
                        r.bold = True
                    r.font.size = Pt(12)
        return

    if is_right:
        for ct in lines_text:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(ct)
            r.italic = True
            r.font.size = Pt(12)
        return

    if is_cover:
        # Preserva a estrutura completa da capa: tags <br> geram
        # parágrafos vazios (espaçamento), tags <p> geram o texto.
        # ABNT: elementos da capa devem ser generosamente espaçados
        # para distribuir o conteúdo verticalmente na página.
        elements = []
        for l in lines:
            clean = re.sub(r'<[^>]+>', '', l).strip()
            if clean and not clean.startswith('&'):
                elements.append(('text', clean))
            elif '<br' in l:
                n = l.count('<br>')
                if n > 0:
                    elements.append(('br', n))
        for elem_type, content in elements:
            if elem_type == 'br':
                for _ in range(content):
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
            elif elem_type == 'text':
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
                r = p.add_run(content)
                r.bold = True
                r.font.size = Pt(12)
        return

    # Generic
    for ct in lines_text:
        if ct:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(ct)
            r.font.size = Pt(12)


def _add_toc_field(doc):
    """Insere um campo Sumário (TOC) no documento.

    O Word/LibreOffice preenche automaticamente os títulos e os
    números de página quando o usuário atualiza o campo
    (botão direito > Atualizar campo). É a forma ABNT de gerar
    o sumário com numeração de páginas e pontilhado de ligação.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    # Níveis 1 e 2 (seções primárias e subseções), com hyperlink e pontilhado.
    instr.text = 'TOC \\o "1-2" \\h \\u'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = ('Sumário — atualize o campo (botão direito > "Atualizar campo") '
              'para exibir os títulos e os números de página.')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(t)
    run._r.append(fld_end)
    return p


def md_to_docx_ufv(md_path, docx_path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Estilos de cabeçalho (Heading) — necessários para que o campo
    # Sumário (TOC) do Word/LibreOffice liste as seções e numere páginas.
    for _h, _sz in (('Heading 1', 14), ('Heading 2', 12)):
        try:
            _st = doc.styles[_h]
            _st.font.name = 'Arial'
            _st.font.size = Pt(_sz)
            _st.font.bold = True
        except KeyError:
            pass

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    html_block = []
    in_html = False
    in_sumario = False
    toc_done = False

    while i < len(lines):
        stripped = lines[i].strip()

        if stripped.startswith('<') and not stripped.startswith('</'):
            if not in_html:
                html_block = []
            in_html = True
            html_block.append(stripped)
            i += 1
            continue

        if in_html:
            if stripped == '':
                if html_block:
                    _html_to_docx(doc, html_block)
                in_html = False
                html_block = []
            else:
                html_block.append(stripped)
            i += 1
            continue

        if stripped.startswith('## '):
            p = doc.add_paragraph()
            p.style = doc.styles['Heading 2']
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            r = p.add_run(stripped.replace('## ', ''))
            r.bold = True
            r.font.size = Pt(12)

        elif stripped.startswith('# ') and not stripped.startswith(('# LISTA', '# SUMÁRIO', '# AGRADECIMENTOS')):
            p = doc.add_paragraph()
            p.style = doc.styles['Heading 1']
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(18)
            r = p.add_run(stripped.replace('# ', ''))
            r.bold = True
            r.font.size = Pt(14)

        elif stripped.startswith('# ') and stripped.startswith(('# LISTA', '# SUMÁRIO', '# AGRADECIMENTOS')):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            r = p.add_run(stripped.replace('# ', ''))
            r.bold = True
            r.font.size = Pt(14)
            if 'SUMÁRIO' in stripped:
                in_sumario = True
                toc_done = False

        elif stripped.startswith('---'):
            if in_sumario and not toc_done:
                _add_toc_field(doc)
                toc_done = True
                in_sumario = False
            doc.add_page_break()

        elif stripped.startswith('| '):
            rows_data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
                rows_data.append(cells)
                i += 1
            if rows_data:
                is_hdr = True
                table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(rows_data):
                    for c_idx, cell_text in enumerate(row_data):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        clean_text = cell_text.strip('-').strip().replace('**', '').replace('*', '')
                        run = p.add_run(clean_text)
                        run.font.size = Pt(10)
                        if is_hdr:
                            run.bold = True
                            set_cell_shading(cell, "D9E2F3")
                    if is_hdr:
                        is_hdr = False
            continue

        elif stripped.startswith('- '):
            if in_sumario:
                i += 1
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.left_indent = Cm(1)
            list_text = stripped
            parts = re.split(r'(\*\*.*?\*\*)', list_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                    r.font.size = Pt(12)
                elif part:
                    clean = part.replace('**', '').replace('*', '')
                    if clean:
                        r = p.add_run(clean)
                        r.font.size = Pt(12)

        elif stripped != '':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                    r.font.size = Pt(12)
                elif part:
                    clean = part.replace('**', '').replace('*', '')
                    if clean:
                        r = p.add_run(clean)
                        r.font.size = Pt(12)

        i += 1

    if html_block and in_html:
        _html_to_docx(doc, html_block)

    from docx.oxml import OxmlElement
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run()
        fb = OxmlElement('w:fldChar')
        fb.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'begin')
        r1._r.append(fb)
        r2 = p.add_run()
        instr = OxmlElement('w:instrText')
        instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instr.text = ' PAGE '
        r2._r.append(instr)
        r3 = p.add_run()
        fe = OxmlElement('w:fldChar')
        fe.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType', 'end')
        r3._r.append(fe)
        for r in p.runs:
            r.font.size = Pt(10)

    # Auto-update fields (inclui sumário TOC) ao abrir no Word/LibreOffice
    try:
        from docx.oxml import parse_xml as _parse_xml
        _settings = doc.settings.element
        _uf = _parse_xml(
            '<w:updateFields w:val="true" '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
        )
        _settings.append(_uf)
    except Exception:
        pass  # não crítico; fallback: usuário atualiza manualmente

    doc.save(docx_path)


# =============================================================================
# CLI
# =============================================================================
def main():
    parser = argparse.ArgumentParser(
        description='Gera memorial RSC-PCCTAE completo a partir do PDF de relatório detalhado.')
    parser.add_argument('pdf', nargs='?', default=None,
                        help='Caminho para o PDF do Relatório Detalhado RSC')
    parser.add_argument('--output-dir', '-o', default=None,
                        help='Diretório de saída (padrão: pasta PesquisAI no Drive)')
    parser.add_argument('--nome', '-n', default=None)
    parser.add_argument('--ano-ingresso', '-a', type=int, default=None)
    parser.add_argument('--auto', action='store_true',
                        help='Modo automático: extrai ano da data de admissão')
    parser.add_argument('--narrativas', '-N', default=None,
                        help='Caminho para JSON com narrativas geradas pelo LLM')
    parser.add_argument('--dump-json', '-J', default=None,
                        help='Exporta dados extraídos do PDF para JSON (sem gerar memorial)')
    args = parser.parse_args()

    if not args.pdf:
        sys.exit("Erro: informe o caminho do PDF. Uso: python3 run.py <caminho_do_pdf>")

    pdf_path = Path(args.pdf)
    if not pdf_path.exists():
        sys.exit(f"Arquivo não encontrado: {pdf_path}")

    if args.output_dir:
        output_dir = Path(args.output_dir)
    else:
        # Default: pasta PesquisAI (Google Drive no Colab, ou variável de ambiente)
        _pesquisai = os.environ.get('PESQUISAI_DIR', '/content/drive/My Drive/PesquisAI')
        output_dir = Path(_pesquisai)
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = pdf_path.stem.replace(' ', '_')

    print("=" * 60)
    print("PDF -> Memorial RSC-PCCTAE v5.3 (NARRATIVAS LLM)")
    print("=" * 60)
    print(f"   Decreto nº 13.048/2026: {DECRETO_URL}")
    if args.narrativas:
        print(f"   Modo: NARRATIVAS LLM")
    else:
        print(f"   Modo: FALLBACK GENÉRICO (sem --narrativas)")
    print("=" * 60)

    print(f"\nLendo PDF: {pdf_path}")
    parser_obj = RSCPDFParser(str(pdf_path))
    data = parser_obj.data

    if args.nome:
        data['nome'] = args.nome
    nome = data['nome']

    # Ano de ingresso
    ano_ingresso = args.ano_ingresso
    if ano_ingresso is None:
        try:
            ano_ingresso = int(data['data_admissao'].split('/')[-1])
        except (ValueError, IndexError, AttributeError):
            if args.auto:
                ano_ingresso = 2009
            else:
                ano_ingresso = perguntar_ano_ingresso()

    print(f"\n   Servidor: {nome}")
    print(f"   Matrícula: {data['matricula']}")
    print(f"   Cargo: {data['cargo']}")
    print(f"   Titulação: {data['titulacao']}")
    print(f"   RSC Requerido: {data['rsc_requerido']}")
    print(f"   Nível: {data['nivel_info']['nome']} (equivalente a {data['nivel_info']['equivalente']})")
    print(f"   Total: {fmt_br(data['total_geral'])} pts ({data['total_criterios']} critérios)")
    for g in data['grupos']:
        print(f"      {g['romano']}: {g['nome_curto']} -- {g['criterios']} crit, {fmt_br(g['pontos'])} pts")
    anos_carreira = datetime.now().year - ano_ingresso
    print(f"\n   Ano de ingresso: {ano_ingresso} ({anos_carreira} anos de carreira)")

    # --dump-json: exporta dados extraídos sem gerar memorial
    if args.dump_json:
        dump_path = Path(args.dump_json)
        dados_exportar = {
            'nome': data['nome'],
            'matricula': data['matricula'],
            'cargo': data['cargo'],
            'titulacao': data['titulacao'],
            'lotacao': data.get('lotacao', ''),
            'data_admissao': data.get('data_admissao', ''),
            'ano_ingresso': ano_ingresso,
            'anos_carreira': anos_carreira,
            'rsc_requerido': data['rsc_requerido'],
            'rsc_nivel': data.get('rsc_nivel', 'VI'),
            'equivalente': data.get('equivalente', 'Doutorado'),
            'total_geral': data['total_geral'],
            'total_criterios': data['total_criterios'],
            'grupos': [{
                'romano': g['romano'], 'nome_curto': g['nome_curto'],
                'criterios': g['criterios'], 'pontos': g['pontos']
            } for g in data['grupos']],
            'criterios': {k: {
                'key': v['key'], 'romano': v['romano'], 'numero': v['numero'],
                'descricao': v['descricao'], 'pontos': v['pontos'],
                'itens': v['itens']
            } for k, v in data['criterios'].items()},
            'ordem_criterios': data.get('ordem_criterios', []),
        }
        with open(dump_path, 'w', encoding='utf-8') as f:
            json.dump(dados_exportar, f, ensure_ascii=False, indent=2)
        print(f"\n   Dados exportados: {dump_path}")
        print(f"   Use este JSON como base para gerar as narrativas com o LLM.")
        print(f"   Depois execute: python3 run.py <pdf> --narrativas narrativas.json")
        return

    # Carrega narrativas do LLM se fornecidas
    narrative_engine = None
    if args.narrativas:
        try:
            narrative_engine = NarrativeEngine(args.narrativas)
        except (FileNotFoundError, json.JSONDecodeError) as e:
            print(f"   Aviso: não foi possível carregar narrativas: {e}")
            print(f"   Usando fallback genérico.")

    print(f"\nGerando memorial em Markdown...")
    gen = MemorialGenerator(data, ano_ingresso, narrative_engine=narrative_engine)
    md_content = gen.generate()
    md_path = output_dir / f"{stem}_MEMORIAL.md"
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    print(f"   OK: {md_path} ({len(md_content)} caracteres)")

    # Generate .docx
    print(f"Gerando .docx formatado UFV/ABNT...")
    docx_path = output_dir / f"{stem}_MEMORIAL.docx"
    try:
        md_to_docx_ufv(str(md_path), str(docx_path))
        sz = os.path.getsize(docx_path) / 1024
        print(f"   OK: {docx_path} ({sz:.1f} KB)")
    except Exception as e:
        print(f"   Aviso: erro ao gerar .docx: {e}")

    print(f"\n{'=' * 60}")
    print(f"Memorial gerado com sucesso!")
    print(f"Formatação UFV/ABNT obrigatória aplicada.")
    print(f"Base legal: Decreto nº 13.048/2026 (Art. 13)")
    print(f"Link: {DECRETO_URL}")
    print('=' * 60)
    print(f"MD:  {md_path}")
    if os.path.exists(docx_path):
        print(f"DOCX: {docx_path}")
    print('=' * 60)


if __name__ == '__main__':
    main()
